import os
import re
import locale
from datetime import datetime
import pandas as pd
from tkinter import Tk, Label, Button, Scrollbar, filedialog, messagebox, StringVar, Canvas
from tkinter.ttk import Progressbar
from docx import Document
from docx.shared import Pt
import threading

class STAutomationApp:
    def __init__(self, root):
        self.root = root
        self.root.title("ST Automation Tool")

        # File paths
        self.excel_path = None
        self.template_path = None
        self.output_dir = None

        # GUI Components
        Label(root, text="ST Automation Tool", font=("Arial", 16)).pack(pady=10)

        Button(root, text="Select Excel File", command=self.select_excel).pack(pady=5)
        self.excel_label = Label(root, text="No file selected", font=("Arial", 10))
        self.excel_label.pack(pady=2)

        Button(root, text="Select Word Template", command=self.select_template).pack(pady=5)
        self.template_label = Label(root, text="No file selected", font=("Arial", 10))
        self.template_label.pack(pady=2)

        Button(root, text="Select Output Directory", command=self.select_output_dir).pack(pady=5)
        self.output_label = Label(root, text="No directory selected", font=("Arial", 10))
        self.output_label.pack(pady=2)

        Button(root, text="Run Automation", command=self.start_automation_thread).pack(pady=10)

        # Progress bar
        self.progress = Progressbar(root, orient="horizontal", length=400, mode="determinate")
        self.progress.pack(pady=10)

        # Loading circle
        self.canvas = Canvas(root, width=50, height=50, highlightthickness=0)
        self.canvas.pack(pady=10)
        self.circle = self.canvas.create_oval(10, 10, 40, 40, outline="blue", width=3)
        self.angle = 0
        self.loading = False

        # Log for results
        self.result_var = StringVar()
        self.result_label = Label(root, textvariable=self.result_var, font=("Arial", 10), justify="left")
        self.result_label.pack(pady=10)

    def select_excel(self):
        self.excel_path = filedialog.askopenfilename(title="Select Excel File", filetypes=[("Excel files", "*.xlsx")])
        if self.excel_path:
            self.excel_label.config(text=f"Selected: {os.path.basename(self.excel_path)}")

    def select_template(self):
        self.template_path = filedialog.askopenfilename(title="Select Word Template", filetypes=[("Word files", "*.docx")])
        if self.template_path:
            self.template_label.config(text=f"Selected: {os.path.basename(self.template_path)}")

    def select_output_dir(self):
        self.output_dir = filedialog.askdirectory(title="Select Output Directory")
        if self.output_dir:
            self.output_label.config(text=f"Selected: {self.output_dir}")

    def format_dates(self, row):
        locale.setlocale(locale.LC_TIME, "id_ID.UTF-8")
        start_date = pd.to_datetime(row['Start Date'])
        end_date = pd.to_datetime(row['End Date'])
        start_day = start_date.strftime("%A")
        start_formatted = start_date.strftime("%d %B %Y")
        end_day = end_date.strftime("%A")
        end_formatted = end_date.strftime("%d %B %Y")
        return f"{start_day} - {end_day}/{start_formatted} - {end_formatted}"

    def format_signatures(self, row):
        locale.setlocale(locale.LC_TIME, "id_ID.UTF-8")
        start_date = pd.to_datetime(row['Start Date'])
        start_formatted = start_date.strftime("%d %B %Y")
        return f"Surabaya, {start_formatted}"

    def replace_surabaya_paragraph(self, doc, row):
        signature_date = self.format_signatures(row)
        pattern = r"Surabaya,\s+\d{1,2}\s+\w+\s+\d{4}"
        for paragraph in doc.paragraphs:
            if re.search(pattern, paragraph.text):
                paragraph.text = re.sub(pattern, signature_date, paragraph.text)

    def sanitize_filename(self, filename):
        return "".join(c for c in filename if c.isalnum() or c in " _-()").strip()

    def run_automation(self):
        if not self.excel_path or not self.template_path or not self.output_dir:
            messagebox.showerror("Error", "Please select all required files and output directory before running.")
            return

        try:
            self.result_var.set("Starting automation process...")
            self.root.update_idletasks()

            data = pd.read_excel(self.excel_path, sheet_name="Individu")
            doc = Document(self.template_path)
            indv_output_folder = os.path.join(self.output_dir, "rektorcup/indv")
            os.makedirs(indv_output_folder, exist_ok=True)

            self.progress["value"] = 0
            self.progress["maximum"] = len(data)

            # result_messages = []

            for index, row in data.iterrows():
                table = doc.tables[1]
                for _ in range(len(table.rows) - 1):
                    table._tbl.remove(table.rows[1]._tr)

                lecturer_name = row['Nama Dosen PA']
                lecturer_nidn = row['NIDN Dosen PA']
                lecturer_row = table.add_row().cells
                lecturer_row[0].text = "1"
                lecturer_row[1].text = lecturer_name
                lecturer_row[2].text = str(lecturer_nidn)
                lecturer_row[3].text = "Dosen Pendamping"

                student_name = row['Name']
                student_nis = row['NIS']
                student_department = row['Major']
                student_row = table.add_row().cells
                student_row[0].text = "2"
                student_row[1].text = student_name
                student_row[2].text = str(student_nis)
                student_row[3].text = student_department

                info_table = doc.tables[2]
                if len(info_table.rows) > 0:
                    tugas_cell = info_table.rows[0].cells[2]
                    tugas_cell.text = f"Mengikuti {row['Activity']}"

                if len(info_table.rows) > 1:
                    haritanggal_cell = info_table.rows[1].cells[2]
                    haritanggal_cell.text = self.format_dates(row)

                self.replace_surabaya_paragraph(doc, row)

                for paragraph in doc.paragraphs:
                    if "No. :" in paragraph.text:
                        paragraph.text = re.sub(r"No\. :\s*[\w/-]+", f"No. : {row['Nomor ST']}", paragraph.text)

                sanitized_name = self.sanitize_filename(row['Name'])
                sanitized_activity = self.sanitize_filename(row['Activity'])
                indv_output_path = f"{indv_output_folder}/ST_{sanitized_name} ({sanitized_activity}).docx"
                doc.save(indv_output_path)

                # result_messages.append(f"Saved: {os.path.basename(indv_output_path)}")

                self.progress["value"] += 1
                self.root.update_idletasks()

            # self.result_var.set("\n".join(result_messages))
            messagebox.showinfo("Process Completed", "Automation process completed successfully.")

        except Exception as e:
            messagebox.showerror("Error", f"An error occurred: {str(e)}")
            self.result_var.set(f"Error: {str(e)}")

        finally:
            self.stop_loading()

    def animate_loading(self):
        while self.loading:
            self.angle = (self.angle + 10) % 360
            x0, y0, x1, y1 = 10, 10, 40, 40
            extent = self.angle
            self.canvas.delete("arc")
            self.canvas.create_arc(x0, y0, x1, y1, start=extent, extent=90, outline="blue", width=3, style="arc", tags="arc")
            self.root.update_idletasks()
            self.root.after(50)

    def start_loading(self):
        self.loading = True
        threading.Thread(target=self.animate_loading, daemon=True).start()

    def stop_loading(self):
        self.loading = False
        self.canvas.delete("arc")

    def start_automation_thread(self):
        self.start_loading()
        threading.Thread(target=self.run_automation, daemon=True).start()

if __name__ == "__main__":
    root = Tk()
    app = STAutomationApp(root)
    root.mainloop()
